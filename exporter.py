"""
exporter.py
-----------
Phase 6 - Professional Export Engine.
Converts Markdown research report into PDF and DOCX.

Dependencies:
    pip install fpdf2 python-docx
"""

import os
import re
from datetime import datetime
from typing import Optional


# ---------------------------------------------------------------------------
# Sanitizer — strips ALL non-latin-1 chars (emoji, unicode) for fpdf2
# ---------------------------------------------------------------------------

def _sanitize(text: str) -> str:
    """Remove every character fpdf2 core fonts cannot render."""
    # Step 1: strip emoji and symbols above latin-1 range
    text = re.sub(r"[^\x00-\xFF]", "", text)
    # Step 2: encode/decode to catch any remaining edge cases
    text = text.encode("latin-1", errors="ignore").decode("latin-1")
    # Step 3: strip any remaining non-printable characters except newline/tab
    text = re.sub(r"[^\x09\x0A\x20-\x7E\xA0-\xFF]", "", text)
    return text


def _strip_md(text: str) -> str:
    """Strip inline markdown AND sanitize for PDF."""
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)  # links → display text
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)            # bold
    text = re.sub(r"\*(.+?)\*",   r"\1", text)              # italic
    text = re.sub(r"`(.+?)`",     r"\1", text)              # inline code
    text = re.sub(r"^#+\s*",      "",    text)              # heading markers
    return _sanitize(text)


# ---------------------------------------------------------------------------
# Markdown tokenizer
# ---------------------------------------------------------------------------

def _parse_markdown(md: str) -> list[dict]:
    tokens = []
    lines  = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.strip().startswith("```"):
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            tokens.append({"type": "code_block", "text": "\n".join(block), "level": None})
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,3})\s+(.+)$", line)
        if m:
            tokens.append({"type": "heading", "text": m.group(2).strip(), "level": len(m.group(1))})
            i += 1
            continue

        # HR
        if re.match(r"^[-*_]{3,}$", line.strip()):
            tokens.append({"type": "hr", "text": "", "level": None})
            i += 1
            continue

        # Bullet
        m = re.match(r"^[\*\-\*\•]\s+(.+)$", line)
        if m:
            tokens.append({"type": "bullet", "text": m.group(1).strip(), "level": None})
            i += 1
            continue

        # Numbered
        m = re.match(r"^\d+[\.\)]\s+(.+)$", line)
        if m:
            tokens.append({"type": "numbered", "text": m.group(1).strip(), "level": None})
            i += 1
            continue

        # Blank
        if not line.strip():
            tokens.append({"type": "blank", "text": "", "level": None})
            i += 1
            continue

        # Paragraph
        tokens.append({"type": "paragraph", "text": line.strip(), "level": None})
        i += 1

    return tokens


# ---------------------------------------------------------------------------
# PDF Exporter
# ---------------------------------------------------------------------------

def export_pdf(
    markdown_text: str,
    topic: str,
    output_path: str,
    model: str = "groq/llama-3.3-70b-versatile",
) -> str:
    from fpdf import FPDF, XPos, YPos

    # Sanitize ALL input strings before they touch fpdf2
    topic_safe = _sanitize(topic)
    model_safe = _sanitize(model)
    date_safe  = datetime.now().strftime("%B %d, %Y")

    # Colours
    NAVY  = (15,  40,  80)
    BLUE  = (37,  99, 235)
    GRAY  = (71,  85, 105)
    LGRAY = (226, 232, 240)
    WHITE = (255, 255, 255)
    BLACK = (15,  23,  42)

    class PDF(FPDF):
        def header(self):
            if self.page_no() > 1:
                self.set_draw_color(*BLUE)
                self.set_line_width(0.5)
                self.line(15, 12, 195, 12)
                self.set_font("Helvetica", "I", 8)
                self.set_text_color(*GRAY)
                self.set_xy(15, 14)
                self.cell(0, 5, topic_safe[:80], align="L")

        def footer(self):
            if self.page_no() > 1:
                self.set_y(-15)
                self.set_draw_color(*LGRAY)
                self.set_line_width(0.3)
                self.line(15, self.get_y(), 195, self.get_y())
                self.set_font("Helvetica", "I", 8)
                self.set_text_color(*GRAY)
                self.cell(0, 8,
                    f"Research Assistant  |  Page {self.page_no() - 1}  |  {model_safe}",
                    align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf = PDF()
    pdf.set_margins(20, 20, 20)
    pdf.set_auto_page_break(auto=True, margin=22)

    # ── Cover page ──────────────────────────────────────────────────────────
    pdf.add_page()

    # Dark header bar
    pdf.set_fill_color(*NAVY)
    pdf.rect(0, 0, 210, 75, style="F")

    pdf.set_font("Helvetica", "B", 13)
    pdf.set_text_color(*WHITE)
    pdf.set_xy(20, 22)
    pdf.cell(0, 8, "RESEARCH ASSISTANT", align="L")

    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(147, 197, 253)
    pdf.set_xy(20, 32)
    pdf.cell(0, 6, f"Powered by {model_safe}", align="L")

    # Blue accent stripe
    pdf.set_fill_color(*BLUE)
    pdf.rect(0, 75, 210, 2, style="F")

    # Topic title
    pdf.set_xy(20, 95)
    pdf.set_font("Helvetica", "B", 22)
    pdf.set_text_color(*BLACK)
    pdf.multi_cell(170, 12, topic_safe, align="L")

    pdf.ln(4)
    pdf.set_font("Helvetica", "", 12)
    pdf.set_text_color(*GRAY)
    pdf.set_x(20)
    pdf.cell(0, 8, "Research Report", align="L", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # Divider
    pdf.ln(8)
    pdf.set_draw_color(*LGRAY)
    pdf.set_line_width(0.3)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(6)

    # Metadata rows
    for label, value in [
        ("Date",   date_safe),
        ("Model",  model_safe),
        ("Type",   "Automated AI Research Report"),
    ]:
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(*NAVY)
        pdf.set_x(20)
        pdf.cell(35, 7, label.upper(), align="L")
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(*BLACK)
        pdf.cell(0, 7, _sanitize(value), align="L",
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # ── Content pages ────────────────────────────────────────────────────────
    pdf.add_page()
    tokens      = _parse_markdown(markdown_text)
    num_counter = 0

    for tok in tokens:
        t    = tok["type"]
        text = _strip_md(tok["text"])   # sanitized on every token

        if t == "blank":
            pdf.ln(3)

        elif t == "hr":
            pdf.set_draw_color(*LGRAY)
            pdf.set_line_width(0.3)
            pdf.line(20, pdf.get_y(), 190, pdf.get_y())
            pdf.ln(4)

        elif t == "heading":
            lvl = tok["level"]
            pdf.ln(3)
            if lvl == 1:
                pdf.set_fill_color(*NAVY)
                pdf.rect(20, pdf.get_y(), 170, 10, style="F")
                pdf.set_font("Helvetica", "B", 13)
                pdf.set_text_color(*WHITE)
                pdf.set_x(23)
                pdf.cell(167, 10, text, align="L",
                         new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.ln(2)
            elif lvl == 2:
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_text_color(*BLUE)
                pdf.set_x(20)
                pdf.cell(0, 8, text, align="L",
                         new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                y = pdf.get_y()
                pdf.set_draw_color(*BLUE)
                pdf.set_line_width(0.4)
                pdf.line(20, y, 100, y)
                pdf.ln(3)
            else:
                pdf.set_font("Helvetica", "B", 10)
                pdf.set_text_color(*NAVY)
                pdf.set_x(20)
                pdf.cell(0, 7, text, align="L",
                         new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        elif t == "bullet":
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(*BLACK)
            pdf.set_x(24)
            pdf.cell(5, 6, "-", align="L")
            pdf.set_x(30)
            pdf.multi_cell(160, 6, text, align="L")
            num_counter = 0

        elif t == "numbered":
            num_counter += 1
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(*BLACK)
            pdf.set_x(24)
            pdf.cell(7, 6, f"{num_counter}.", align="L")
            pdf.set_x(32)
            pdf.multi_cell(158, 6, text, align="L")

        elif t == "paragraph":
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(*BLACK)
            pdf.set_x(20)
            pdf.multi_cell(170, 6, text, align="L")
            pdf.ln(2)
            num_counter = 0

        elif t == "code_block":
            pdf.set_fill_color(241, 245, 249)
            pdf.set_font("Courier", "", 8)
            pdf.set_text_color(*GRAY)
            pdf.set_x(20)
            code_text = _sanitize(tok["text"])
            pdf.multi_cell(170, 5, code_text, align="L", fill=True)
            pdf.ln(2)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    pdf.output(output_path)
    return os.path.abspath(output_path)


# ---------------------------------------------------------------------------
# DOCX Exporter
# ---------------------------------------------------------------------------

def export_docx(
    markdown_text: str,
    topic: str,
    output_path: str,
    model: str = "groq/llama-3.3-70b-versatile",
) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width   = Inches(8.5)
    section.page_height  = Inches(11)
    section.left_margin  = section.right_margin  = Inches(1.2)
    section.top_margin   = section.bottom_margin = Inches(1.0)

    def _add_inline(para, text: str):
        """Parse inline markdown → styled runs. Safe against hyperlink crashes."""
        pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^\)]+\))"
        parts   = re.split(pattern, text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = para.add_run(part[1:-1])
                run.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = para.add_run(part[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(71, 85, 105)
            else:
                # Handle markdown links safely — no XML manipulation
                m = re.match(r"\[([^\]]+)\]\(([^\)]+)\)", part)
                if m:
                    display, url = m.group(1), m.group(2)
                    run = para.add_run(f"{display}")
                    run.font.color.rgb = RGBColor(37, 99, 235)
                    run.underline = True
                else:
                    para.add_run(part)

    # ── Cover page ───────────────────────────────────────────────────────────
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cover.paragraph_format.space_before = Pt(60)
    run = cover.add_run(topic)
    run.font.name  = "Calibri"
    run.font.size  = Pt(26)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(15, 40, 80)

    sub = doc.add_paragraph("Research Report")
    sub.runs[0].font.size      = Pt(13)
    sub.runs[0].font.color.rgb = RGBColor(37, 99, 235)
    sub.runs[0].font.name      = "Calibri"

    doc.add_paragraph("")

    for label, value in [
        ("Date",  datetime.now().strftime("%B %d, %Y")),
        ("Model", model),
        ("Type",  "Automated AI Research Report"),
    ]:
        p  = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{label}:  ")
        r1.bold = True
        r1.font.size      = Pt(10)
        r1.font.color.rgb = RGBColor(15, 40, 80)
        r2 = p.add_run(value)
        r2.font.size      = Pt(10)
        r2.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_page_break()

    # ── Content ───────────────────────────────────────────────────────────────
    tokens      = _parse_markdown(markdown_text)
    num_counter = 0

    for tok in tokens:
        t    = tok["type"]
        text = tok["text"]

        if t == "blank":
            continue

        elif t == "heading":
            lvl = min(tok["level"], 3)
            h   = doc.add_heading(
                re.sub(r"[^\x00-\xFF]", "", _strip_md(text)), level=lvl
            )
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT

        elif t == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, text)
            num_counter = 0

        elif t == "numbered":
            num_counter += 1
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, text)

        elif t == "paragraph":
            p = doc.add_paragraph(style="Normal")
            _add_inline(p, text)
            num_counter = 0

        elif t == "code_block":
            p   = doc.add_paragraph()
            run = p.add_run(tok["text"])
            run.font.name      = "Courier New"
            run.font.size      = Pt(9)
            run.font.color.rgb = RGBColor(71, 85, 105)
            p.paragraph_format.left_indent = Inches(0.3)

        elif t == "hr":
            p   = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot  = OxmlElement("w:bottom")
            bot.set(qn("w:val"),   "single")
            bot.set(qn("w:sz"),    "6")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), "2563EB")
            pBdr.append(bot)
            pPr.append(pBdr)

    # Footer with page numbers
    footer      = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.clear()
    run = footer_para.add_run("Research Assistant  |  Page ")
    run.font.size      = Pt(8)
    run.font.color.rgb = RGBColor(148, 163, 184)

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run2 = footer_para.add_run()
    run2.font.size      = Pt(8)
    run2.font.color.rgb = RGBColor(148, 163, 184)
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    return os.path.abspath(output_path)


# ---------------------------------------------------------------------------
# Export all formats
# ---------------------------------------------------------------------------

def export_all(
    markdown_text: str,
    topic: str,
    base_path: str,
    model: str = "groq/llama-3.3-70b-versatile",
) -> dict:
    results = {}
    md_path = base_path if base_path.endswith(".md") else base_path + ".md"
    if os.path.exists(md_path):
        results["md"] = md_path
    stem = md_path.replace(".md", "")

    try:
        pdf_path = stem + ".pdf"
        export_pdf(markdown_text, topic, pdf_path, model)
        results["pdf"] = pdf_path
    except Exception as exc:
        results["pdf_error"] = str(exc)

    try:
        docx_path = stem + ".docx"
        export_docx(markdown_text, topic, docx_path, model)
        results["docx"] = docx_path
    except Exception as exc:
        results["docx_error"] = str(exc)

    return results
